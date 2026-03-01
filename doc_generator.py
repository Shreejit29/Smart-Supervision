from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy


# ---------------------------------------------------
# Utility: Apply Times New Roman Font (Size 14)
# ---------------------------------------------------
def set_font(cell, bold=False):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.bold = bold

            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


# ---------------------------------------------------
# Utility: Add Full Borders to Table
# ---------------------------------------------------
def add_table_borders(table):
    tbl = table._tbl

    for row in tbl.tr_lst:
        for cell in row.tc_lst:
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')

            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '8')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)

            tcPr.append(tcBorders)


# ===================================================
# UPLOAD MODE (Institutional Matrix Excel)
# ===================================================
def generate_individual_doc(faculty, schedule, analysis, template_file=None):

    if template_file is not None:
        doc = Document(template_file)
    else:
        doc = Document()

        heading = doc.add_heading("Individual Supervision Chart", level=1)
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

        p1 = doc.add_paragraph(f"Name: {faculty['name']}")
        p2 = doc.add_paragraph(f"Department: {faculty['department']}")

        for paragraph in [p1, p2]:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)

    # Replace Template Placeholders
    for paragraph in doc.paragraphs:

        if "{{NAME}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{NAME}}", str(faculty["name"]))

        if "{{DEPARTMENT}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{DEPARTMENT}}", str(faculty["department"]))

        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

    # Add Supervision Table
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Sr.No", "Date", "Session", "Time"]

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        set_font(table.rows[0].cells[i], bold=True)

    sr = 1
    first_data_col = min(analysis["date_columns"].keys())

    for date, sessions in schedule.items():
        for session_name, session_time, col_idx in sessions:

            data_index = col_idx - first_data_col

            if data_index < len(faculty["data"]) and faculty["data"][data_index] == 1:

                row_cells = table.add_row().cells
                row_cells[0].text = str(sr)
                row_cells[1].text = date
                row_cells[2].text = session_name
                row_cells[3].text = session_time

                for cell in row_cells:
                    set_font(cell)

                sr += 1

    add_table_borders(table)

    return doc


# ===================================================
# AUTO MODE (Flat Master DataFrame)
# ===================================================
def generate_individual_doc_auto(faculty_name, faculty_department, master_df, template_file=None):

    if template_file is not None:
        doc = Document(template_file)
    else:
        doc = Document()

        heading = doc.add_heading("Individual Supervision Chart", level=1)
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

        p1 = doc.add_paragraph(f"Name: {faculty_name}")
        p2 = doc.add_paragraph(f"Department: {faculty_department}")

        for paragraph in [p1, p2]:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)

    # Replace placeholders
    for paragraph in doc.paragraphs:

        if "{{NAME}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{NAME}}", str(faculty_name))

        if "{{DEPARTMENT}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{DEPARTMENT}}", str(faculty_department))

        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

    faculty_schedule = master_df[
        master_df["Name of faculty"] == faculty_name
    ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Sr.No", "Date", "Session", "Time"]

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        set_font(table.rows[0].cells[i], bold=True)

    sr = 1

    for _, row in faculty_schedule.iterrows():

        row_cells = table.add_row().cells
        row_cells[0].text = str(sr)
        row_cells[1].text = str(row["Date"])
        row_cells[2].text = str(row["Session"])
        row_cells[3].text = str(row["Time"])

        for cell in row_cells:
            set_font(cell)

        sr += 1

    add_table_borders(table)

    return doc


# ===================================================
# COMBINE DOCUMENTS
# ===================================================
def combine_documents(individual_docs):

    master_doc = Document()
    first = True

    for _, doc in individual_docs:

        if not first:
            master_doc.add_page_break()

        first = False

        for element in doc.element.body:
            master_doc.element.body.append(deepcopy(element))

    return master_doc
